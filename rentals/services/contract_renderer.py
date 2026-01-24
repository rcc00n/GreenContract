from io import BytesIO
import logging
import re
from decimal import Decimal
from typing import Iterable

from django.template import engines
from django.utils import timezone
from docx import Document
from pypdf import PdfReader, PdfWriter
from pypdf.generic import BooleanObject, NameObject
from xhtml2pdf import pisa

from ..models import ContractTemplate, Rental

logger = logging.getLogger(__name__)

# Readable placeholder descriptions for the contract template constructor UI.
PLACEHOLDER_GUIDE = [
    (
        "Клиент",
        {
            "customer.full_name": "ФИО клиента",
            "customer.birth_date": "Дата рождения",
            "customer.phone": "Телефон",
            "customer.email": "Эл. почта",
            "customer.license_number": "Водительское удостоверение",
            "customer.license_issued_by": "В.у. выдано",
            "customer.driving_since": "Стаж с",
            "customer.discount_percent": "Скидка, %",
            "customer.passport_series": "Серия паспорта",
            "customer.passport_number": "Номер паспорта",
            "customer.passport_issue_date": "Дата выдачи паспорта",
            "customer.passport_issued_by": "Кем выдан паспорт",
            "customer.registration_address": "Адрес регистрации",
        },
    ),
    (
        "Второй водитель",
        {
            "second_driver.full_name": "ФИО второго водителя",
            "second_driver.birth_date": "Дата рождения",
            "second_driver.phone": "Телефон",
            "second_driver.email": "Эл. почта",
            "second_driver.license_number": "Водительское удостоверение",
            "second_driver.license_issued_by": "В.у. выдано",
            "second_driver.driving_since": "Стаж с",
            "second_driver.discount_percent": "Скидка, %",
            "second_driver.passport_series": "Серия паспорта",
            "second_driver.passport_number": "Номер паспорта",
            "second_driver.passport_issue_date": "Дата выдачи паспорта",
            "second_driver.passport_issued_by": "Кем выдан паспорт",
            "second_driver.registration_address": "Адрес регистрации",
        },
    ),
    (
        "Авто",
        {
            "car.plate_number": "Госномер",
            "car.make": "Марка",
            "car.model": "Модель",
            "car.year": "Год",
            "car.vin": "ВИН / номер кузова",
            "car.sts_number": "Номер СТС",
            "car.sts_issue_date": "Дата выдачи СТС",
            "car.sts_issued_by": "Кем выдано СТС",
            "car.label": "Госномер + марка",
            "car.security_deposit": "Залог",
            "car.security_deposit_text": "Залог прописью",
        },
    ),
    (
        "Договор",
        {
            "rental.contract_number": "Номер договора",
            "rental.start_date": "Дата начала",
            "rental.end_date": "Дата окончания",
            "rental.start_time": "Время начала",
            "rental.end_time": "Время завершения",
            "rental.date_range": "Диапазон дат",
            "rental.duration_days": "Длительность, дней",
            "rental.daily_rate": "Суточная ставка",
            "rental.total_price": "Стоимость",
            "rental.balance_due": "К оплате после предоплаты",
            "rental.balance_due_text": "К оплате прописью",
            "rental.prepayment": "Предоплата",
            "rental.advance_payment_text": "Предоплата прописью",
            "rental.discount_amount": "Скидка суммой",
            "rental.discount_percent": "Скидка, %",
            "rental.car_wash_fee": "Мойка",
            "rental.night_fee_start": "Ночной выход (выдача)",
            "rental.night_fee_end": "Ночной выход (возврат)",
            "rental.delivery_issue_city": "Город выдачи (доставка)",
            "rental.delivery_issue_fee": "Сумма выдачи (доставка)",
            "rental.delivery_return_city": "Город возврата (доставка)",
            "rental.delivery_return_fee": "Сумма возврата (доставка)",
            "rental.operation_regions": "Территория эксплуатации",
            "rental.mileage_limit_km": "Ограничение пробега, км",
            "rental.child_seat_included": "Детское кресло включено",
            "rental.child_seat_count": "Количество кресел",
            "rental.booster_included": "Бустер включен",
            "rental.booster_count": "Количество бустеров",
            "rental.ski_rack_included": "Багажник для лыж включен",
            "rental.ski_rack_count": "Количество багажников",
            "rental.roof_box_included": "Бокс на крышу включен",
            "rental.roof_box_count": "Количество боксов",
            "rental.crossbars_included": "Поперечины включены",
            "rental.crossbars_count": "Количество поперечин",
            "rental.equipment_manual_total": "Фиксированная сумма за оборудование",
            "rental.deal_name": "Удобочитаемое имя сделки",
            "rental.created_via_wizard": "Создано мастером",
        },
    ),
    (
        "Служебное",
        {
            "meta.today": "Текущая дата",
            "meta.generated_at": "Дата и время генерации",
        },
    ),
]

# Legacy Russian placeholder aliases kept for backward compatibility.
PLACEHOLDER_ALIASES_RU = {
    "customer.full_name": "клиент.фио",
    "customer.birth_date": "клиент.дата_рождения",
    "customer.phone": "клиент.телефон",
    "customer.email": "клиент.эл_почта",
    "customer.license_number": "клиент.номер_ву",
    "customer.license_issued_by": "клиент.кем_выдано_ву",
    "customer.driving_since": "клиент.стаж_с",
    "customer.discount_percent": "клиент.скидка_процент",
    "customer.passport_series": "клиент.паспорт_серия",
    "customer.passport_number": "клиент.паспорт_номер",
    "customer.passport_issue_date": "клиент.паспорт_дата_выдачи",
    "customer.passport_issued_by": "клиент.паспорт_кем_выдан",
    "customer.registration_address": "клиент.адрес_регистрации",
    "customer.address": "клиент.адрес",
    "customer.residence_address": "клиент.адрес_проживания",
    "second_driver.full_name": "второй_водитель.фио",
    "second_driver.birth_date": "второй_водитель.дата_рождения",
    "second_driver.phone": "второй_водитель.телефон",
    "second_driver.email": "второй_водитель.эл_почта",
    "second_driver.license_number": "второй_водитель.номер_ву",
    "second_driver.license_issued_by": "второй_водитель.кем_выдано_ву",
    "second_driver.driving_since": "второй_водитель.стаж_с",
    "second_driver.discount_percent": "второй_водитель.скидка_процент",
    "second_driver.passport_series": "второй_водитель.паспорт_серия",
    "second_driver.passport_number": "второй_водитель.паспорт_номер",
    "second_driver.passport_issue_date": "второй_водитель.паспорт_дата_выдачи",
    "second_driver.passport_issued_by": "второй_водитель.паспорт_кем_выдан",
    "second_driver.registration_address": "второй_водитель.адрес_регистрации",
    "second_driver.address": "второй_водитель.адрес",
    "second_driver.residence_address": "второй_водитель.адрес_проживания",
    "car.plate_number": "авто.госномер",
    "car.make": "авто.марка",
    "car.model": "авто.модель",
    "car.year": "авто.год",
    "car.vin": "авто.вин",
    "car.sts_number": "авто.стс_номер",
    "car.sts_issue_date": "авто.стс_дата_выдачи",
    "car.sts_issued_by": "авто.стс_кем_выдано",
    "car.label": "авто.название",
    "car.security_deposit": "авто.залог",
    "car.security_deposit_text": "авто.залог_прописью",
    "rental.contract_number": "аренда.номер_договора",
    "rental.start_date": "аренда.дата_начала",
    "rental.end_date": "аренда.дата_окончания",
    "rental.start_time": "аренда.время_начала",
    "rental.end_time": "аренда.время_окончания",
    "rental.date_range": "аренда.период",
    "rental.duration_days": "аренда.дней",
    "rental.daily_rate": "аренда.суточный_тариф",
    "rental.total_price": "аренда.итоговая_сумма",
    "rental.balance_due": "аренда.к_оплате",
    "rental.balance_due_text": "аренда.к_оплате_прописью",
    "rental.prepayment": "аренда.предоплата",
    "rental.advance_payment_text": "аренда.предоплата_прописью",
    "rental.discount_amount": "аренда.скидка_сумма",
    "rental.discount_percent": "аренда.скидка_процент",
    "rental.airport_fee_start": "аренда.аэропорт_выдача",
    "rental.airport_fee_end": "аренда.аэропорт_возврат",
    "rental.car_wash_fee": "аренда.мойка",
    "rental.night_fee_start": "аренда.ночной_выход_выдача",
    "rental.night_fee_end": "аренда.ночной_выход_возврат",
    "rental.delivery_issue_city": "аренда.доставка_город_выдачи",
    "rental.delivery_issue_fee": "аренда.доставка_стоимость_выдачи",
    "rental.delivery_return_city": "аренда.доставка_город_возврата",
    "rental.delivery_return_fee": "аренда.доставка_стоимость_возврата",
    "rental.operation_regions": "аренда.территория_эксплуатации",
    "rental.mileage_limit_km": "аренда.лимит_пробега_км",
    "rental.child_seat_included": "аренда.кресло_включено",
    "rental.child_seat_count": "аренда.кресло_количество",
    "rental.booster_included": "аренда.бустер_включен",
    "rental.booster_count": "аренда.бустер_количество",
    "rental.ski_rack_included": "аренда.крепления_лыжи_включены",
    "rental.ski_rack_count": "аренда.крепления_лыжи_количество",
    "rental.roof_box_included": "аренда.автобокс_включен",
    "rental.roof_box_count": "аренда.автобокс_количество",
    "rental.crossbars_included": "аренда.поперечины_включены",
    "rental.crossbars_count": "аренда.поперечины_количество",
    "rental.equipment_manual_total": "аренда.оборудование_сумма",
    "rental.deal_name": "аренда.имя_сделки",
    "rental.created_via_wizard": "аренда.создано_мастером",
    "meta.today": "мета.сегодня",
    "meta.generated_at": "мета.дата_генерации",
}


def _build_ru_alias_map() -> dict[str, dict[str, str]]:
    aliases: dict[str, dict[str, str]] = {}
    for eng_key, ru_key in PLACEHOLDER_ALIASES_RU.items():
        try:
            eng_group, eng_attr = eng_key.split(".", 1)
            ru_group, ru_attr = ru_key.split(".", 1)
        except ValueError:
            continue
        group_aliases = aliases.setdefault(ru_group, {})
        group_aliases[ru_attr] = eng_attr
    return aliases


RU_ALIAS_GROUPS = _build_ru_alias_map()


class _BoolDisplay:
    def __init__(self, value: bool):
        self._value = bool(value)

    def __bool__(self):
        return self._value

    def __str__(self):
        return "Да" if self._value else "Нет"

    def __eq__(self, other):
        if isinstance(other, _BoolDisplay):
            return self._value == other._value
        return self._value == other


class _DateFormattingProxy:
    def __init__(
        self,
        obj,
        *,
        date_fields: Iterable[str] = (),
        time_fields: Iterable[str] = (),
        datetime_fields: Iterable[str] = (),
    ):
        self._obj = obj
        self._date_fields = set(date_fields)
        self._time_fields = set(time_fields)
        self._datetime_fields = set(datetime_fields)

    def __getattr__(self, name):
        value = getattr(self._obj, name)
        if name in self._date_fields:
            return _fmt_date(value)
        if name in self._time_fields:
            return _fmt_time(value)
        if name in self._datetime_fields:
            return _fmt_datetime(value)
        return value

    def __str__(self):
        return str(self._obj)


class _EmptyProxy:
    def __init__(self, label: str = ""):
        self._label = label

    def __getattr__(self, name):
        return ""

    def __str__(self):
        return self._label


class _AliasProxy:
    def __init__(self, obj, aliases: dict[str, str] | None = None):
        self._obj = obj
        self._aliases = aliases or {}

    def __getattr__(self, name):
        translated = self._aliases.get(name, name)
        return getattr(self._obj, translated)

    def __str__(self):
        return str(self._obj)


class _RentalTemplateProxy:
    def __init__(self, rental: Rental):
        self._rental = rental
        self.child_seat_included = _BoolDisplay(rental.child_seat_included)
        self.booster_included = _BoolDisplay(rental.booster_included)
        self.ski_rack_included = _BoolDisplay(rental.ski_rack_included)
        self.roof_box_included = _BoolDisplay(rental.roof_box_included)
        self.crossbars_included = _BoolDisplay(rental.crossbars_included)
        self.created_via_wizard = _BoolDisplay(rental.created_via_wizard)
        self.car = _DateFormattingProxy(rental.car, date_fields={"sts_issue_date"})
        self.customer = _DateFormattingProxy(
            rental.customer,
            date_fields={"birth_date", "driving_since", "passport_issue_date"},
        )
        self.second_driver = (
            _DateFormattingProxy(
                rental.second_driver,
                date_fields={"birth_date", "driving_since", "passport_issue_date"},
            )
            if rental.second_driver
            else _EmptyProxy()
        )

    def __getattr__(self, name):
        if name in {"start_date", "end_date"}:
            return _fmt_date(getattr(self._rental, name))
        if name in {"start_time", "end_time"}:
            return _fmt_time(getattr(self._rental, name))
        if name in {"delivery_issue_city", "delivery_return_city"}:
            return _normalize_delivery_city(getattr(self._rental, name, ""))
        if name == "date_range":
            return _format_date_range(self._rental.start_date, self._rental.end_date)
        if name == "deal_name":
            return _format_deal_name(self._rental)
        return getattr(self._rental, name)


def get_contract_context(rental: Rental) -> dict:
    start_date = rental.start_date
    end_date = rental.end_date
    duration_days = rental.duration_days if start_date and end_date else None
    date_range = _format_date_range(start_date, end_date)
    rental_proxy = _RentalTemplateProxy(rental)
    meta = {
        "generated_at": _fmt_datetime(timezone.localtime()),
        "today": _fmt_date(timezone.localdate()),
    }
    meta_ru = {**meta}
    meta_ru.update(
        {ru_key: meta.get(eng_key, "") for ru_key, eng_key in RU_ALIAS_GROUPS.get("мета", {}).items()}
    )
    return {
        "rental": rental_proxy,
        "car": rental_proxy.car,
        "customer": rental_proxy.customer,
        "second_driver": rental_proxy.second_driver,
        "аренда": _AliasProxy(rental_proxy, RU_ALIAS_GROUPS.get("аренда", {})),
        "авто": _AliasProxy(rental_proxy.car, RU_ALIAS_GROUPS.get("авто", {})),
        "клиент": _AliasProxy(rental_proxy.customer, RU_ALIAS_GROUPS.get("клиент", {})),
        "второй_водитель": _AliasProxy(rental_proxy.second_driver, RU_ALIAS_GROUPS.get("второй_водитель", {})),
        "rental_duration_days": duration_days,
        "rental_date_range": date_range,
        "meta": meta,
        "мета": meta_ru,
    }


def _normalize_html_charset(html: str, target: str = "utf-8") -> str:
    """
    Force HTML to declare UTF-8 so browsers decode Russian text correctly.

    Word-exported HTML often keeps a windows-1251 meta tag while Django still
    encodes the response as UTF-8, which results in mojibake. We rewrite the
    charset declaration (both <meta charset=...> and http-equiv variants) or
    inject one if missing.
    """
    charset_re = re.compile(r'(<meta[^>]+charset=)([\"\\\']?)([^\"\\\' >]+)', re.IGNORECASE)
    http_equiv_re = re.compile(
        r'(<meta[^>]+http-equiv=[\"\\\']?content-type[\"\\\']?[^>]+charset=)([^\"\\\' >]+)',
        re.IGNORECASE,
    )

    updated, replaced = charset_re.subn(rf"\1\2{target}", html, count=1)
    if replaced == 0:
        updated, replaced = http_equiv_re.subn(rf"\1{target}", html, count=1)

    if replaced == 0:
        head_match = re.search(r"<head[^>]*>", html, flags=re.IGNORECASE)
        meta_tag = f'<meta charset="{target}">'
        if head_match:
            insert_at = head_match.end()
            updated = html[:insert_at] + meta_tag + html[insert_at:]
        else:
            updated = meta_tag + html

    return updated


def _fmt_date(value) -> str:
    return value.strftime("%d.%m.%Y") if value else ""


def _fmt_datetime(value) -> str:
    return value.strftime("%d.%m.%Y %H:%M") if value else ""


def _fmt_time(value) -> str:
    return value.strftime("%H:%M") if value else ""


def _format_date_range(start_date, end_date) -> str:
    if not start_date and not end_date:
        return ""
    start = _fmt_date(start_date)
    end = _fmt_date(end_date)
    if start and end:
        return f"{start} — {end}"
    return start or end


def _normalize_delivery_city(value: str | None) -> str:
    if not value:
        return ""
    text = str(value).strip()
    return re.sub(r"\s*[-–—]\s*\d+\s*$", "", text)


def _format_deal_name(rental: Rental) -> str:
    contract = rental.contract_number or "-----"
    last_name = rental.customer_last_name or "—"
    car_piece = ""
    if rental.car_id:
        car_piece = f"{rental.car.plate_number} {rental.car.make}".strip()
    date_piece = _fmt_date(rental.start_date) if rental.start_date else ""
    return f"{contract}/{last_name}/{car_piece}/{date_piece}"


def _fmt_decimal(value) -> str:
    if value is None:
        return ""
    try:
        number = Decimal(value)
        return f"{number:.2f}".rstrip("0").rstrip(".")
    except Exception:
        return str(value)


def _fmt_bool(value) -> str:
    return "Да" if bool(value) else "Нет"


def build_placeholder_values(rental: Rental) -> dict[str, str]:
    """Flatten rental/car/customer data into string placeholders."""
    customer = rental.customer
    second_driver = rental.second_driver
    car = rental.car
    start_date = rental.start_date
    end_date = rental.end_date
    duration_days = rental.duration_days if start_date and end_date else None
    date_range = _format_date_range(start_date, end_date)
    address_primary = customer.registration_address or customer.address or customer.residence_address
    address_second = ""
    if second_driver:
        address_second = (
            second_driver.registration_address
            or second_driver.address
            or second_driver.residence_address
        )

    values = {
        # Client
        "customer.full_name": customer.full_name,
        "customer.birth_date": _fmt_date(customer.birth_date),
        "customer.phone": customer.phone,
        "customer.email": customer.email,
        "customer.license_number": customer.license_number,
        "customer.license_issued_by": customer.license_issued_by,
        "customer.driving_since": _fmt_date(customer.driving_since),
        "customer.discount_percent": _fmt_decimal(customer.discount_percent),
        "customer.passport_series": customer.passport_series,
        "customer.passport_number": customer.passport_number,
        "customer.passport_issue_date": _fmt_date(customer.passport_issue_date),
        "customer.passport_issued_by": customer.passport_issued_by,
        "customer.address": address_primary,
        "customer.registration_address": address_primary,
        "customer.residence_address": address_primary,
        # Second driver
        "second_driver.full_name": second_driver.full_name if second_driver else "",
        "second_driver.birth_date": _fmt_date(second_driver.birth_date) if second_driver else "",
        "second_driver.phone": second_driver.phone if second_driver else "",
        "second_driver.email": second_driver.email if second_driver else "",
        "second_driver.license_number": second_driver.license_number if second_driver else "",
        "second_driver.license_issued_by": second_driver.license_issued_by if second_driver else "",
        "second_driver.driving_since": _fmt_date(second_driver.driving_since) if second_driver else "",
        "second_driver.discount_percent": _fmt_decimal(second_driver.discount_percent) if second_driver else "",
        "second_driver.passport_series": second_driver.passport_series if second_driver else "",
        "second_driver.passport_number": second_driver.passport_number if second_driver else "",
        "second_driver.passport_issue_date": _fmt_date(second_driver.passport_issue_date) if second_driver else "",
        "second_driver.passport_issued_by": second_driver.passport_issued_by if second_driver else "",
        "second_driver.address": address_second,
        "second_driver.registration_address": address_second,
        "second_driver.residence_address": address_second,
        # Car
        "car.plate_number": car.plate_number,
        "car.make": car.make,
        "car.model": car.model,
        "car.year": car.year,
        "car.vin": car.vin,
        "car.sts_number": car.sts_number,
        "car.sts_issue_date": _fmt_date(car.sts_issue_date),
        "car.sts_issued_by": car.sts_issued_by,
        "car.label": str(car),
        "car.security_deposit": _fmt_decimal(car.security_deposit),
        "car.security_deposit_text": car.security_deposit_text,
        # Rental
        "rental.contract_number": rental.contract_number or "",
        "rental.start_date": _fmt_date(start_date),
        "rental.end_date": _fmt_date(end_date),
        "rental.start_time": _fmt_time(rental.start_time),
        "rental.end_time": _fmt_time(rental.end_time),
        "rental.date_range": date_range,
        "rental.duration_days": duration_days or "",
        "rental.daily_rate": _fmt_decimal(rental.daily_rate),
        "rental.total_price": _fmt_decimal(rental.total_price),
        "rental.balance_due": _fmt_decimal(rental.balance_due),
        "rental.balance_due_text": rental.balance_due_text,
        "rental.prepayment": _fmt_decimal(rental.prepayment),
        "rental.advance_payment_text": rental.advance_payment_text,
        "rental.discount_amount": _fmt_decimal(rental.discount_amount),
        "rental.discount_percent": _fmt_decimal(rental.discount_percent),
        "rental.airport_fee_start": _fmt_decimal(getattr(rental, "airport_fee_start", "")),
        "rental.airport_fee_end": _fmt_decimal(getattr(rental, "airport_fee_end", "")),
        "rental.car_wash_fee": _fmt_decimal(getattr(rental, "car_wash_fee", "")),
        "rental.night_fee_start": _fmt_decimal(rental.night_fee_start),
        "rental.night_fee_end": _fmt_decimal(rental.night_fee_end),
        "rental.delivery_issue_city": _normalize_delivery_city(rental.delivery_issue_city),
        "rental.delivery_issue_fee": _fmt_decimal(rental.delivery_issue_fee),
        "rental.delivery_return_city": _normalize_delivery_city(rental.delivery_return_city),
        "rental.delivery_return_fee": _fmt_decimal(rental.delivery_return_fee),
        "rental.operation_regions": rental.operation_regions,
        "rental.mileage_limit_km": _fmt_decimal(rental.mileage_limit_km),
        "rental.child_seat_included": _fmt_bool(rental.child_seat_included),
        "rental.child_seat_count": rental.child_seat_count,
        "rental.booster_included": _fmt_bool(rental.booster_included),
        "rental.booster_count": rental.booster_count,
        "rental.ski_rack_included": _fmt_bool(rental.ski_rack_included),
        "rental.ski_rack_count": rental.ski_rack_count,
        "rental.roof_box_included": _fmt_bool(rental.roof_box_included),
        "rental.roof_box_count": rental.roof_box_count,
        "rental.crossbars_included": _fmt_bool(rental.crossbars_included),
        "rental.crossbars_count": rental.crossbars_count,
        "rental.equipment_manual_total": _fmt_decimal(rental.equipment_manual_total),
        "rental.deal_name": rental.deal_name,
        # Meta
        "meta.today": _fmt_date(timezone.localdate()),
        "meta.generated_at": _fmt_datetime(timezone.localtime()),
    }

    for eng_key, ru_key in PLACEHOLDER_ALIASES_RU.items():
        if eng_key in values and ru_key not in values:
            values[ru_key] = values[eng_key]

    return {key: "" if value is None else str(value) for key, value in values.items()}


def placeholder_token_map(rental: Rental) -> dict[str, str]:
    """
    Build a mapping of popular placeholder token variants so DOCX/PDF can
    perform quick replacements.
    """
    mapping: dict[str, str] = {}
    values = build_placeholder_values(rental)
    for dotted, value in values.items():
        flat = dotted.replace(".", "_")
        variants = (
            f"{{{{ {dotted} }}}}",
            f"{{{{{dotted}}}}}",
            f"{{{{ {flat} }}}}",
            f"{{{{{flat}}}}}",
            flat,
        )
        for token in variants:
            mapping[token] = value
    return mapping


def placeholder_guide() -> list[dict]:
    """Provide grouped placeholder hints for the constructor UI."""
    groups = []
    for title, items in PLACEHOLDER_GUIDE:
        groups.append(
            {
                "title": title,
                "items": [
                    {
                        "token": f"{{{{ {key} }}}}",
                        "alt": key.replace(".", "_"),
                        "description": description,
                    }
                    for key, description in items.items()
                ],
            }
        )
    return groups


def render_html_template(contract_template: ContractTemplate, rental: Rental) -> str:
    if not contract_template.body_html:
        raise ValueError("Разметка веб-шаблона пуста.")

    django_engine = engines["django"]
    template = django_engine.from_string(contract_template.body_html)
    context = get_contract_context(rental)
    html = template.render(context)
    return _normalize_html_charset(html)


def render_html_to_pdf(html: str) -> bytes:
    """Convert HTML to PDF bytes using xhtml2pdf."""
    output = BytesIO()
    result = pisa.CreatePDF(html, dest=output, encoding="utf-8")
    output.seek(0)
    if result.err:
        raise ValueError("Не удалось сформировать ПДФ из веб-шаблона.")
    return output.getvalue()


def _replace_in_paragraphs(paragraphs: Iterable, mapping: dict[str, str]):
    """Replace placeholders in a list of DOCX paragraphs."""
    for paragraph in paragraphs:
        original = paragraph.text
        updated = original
        for key, value in mapping.items():
            if key in updated:
                updated = updated.replace(key, value)
        if updated != original:
            paragraph.text = updated


def render_docx(contract_template: ContractTemplate, rental: Rental) -> BytesIO:
    """
    Загружает DOCX-шаблон и заменяет плейсхолдеры вида {{ customer.full_name }}.
    Применяет замену в абзацах, таблицах, колонтитулах и нижних колонтитулах.
    """
    document = Document(contract_template.file.path)
    mapping = placeholder_token_map(rental)

    _replace_in_paragraphs(document.paragraphs, mapping)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs, mapping)

    for section in document.sections:
        _replace_in_paragraphs(section.header.paragraphs, mapping)
        _replace_in_paragraphs(section.footer.paragraphs, mapping)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


def render_pdf(contract_template: ContractTemplate, rental: Rental) -> BytesIO:
    """
    Render a PDF contract from either HTML body (converted to PDF) or a
    fillable PDF template with AcroForm fields.
    """
    if contract_template.file:
        return _fill_pdf_form(contract_template.file.path, rental)

    if contract_template.body_html:
        html = render_html_template(contract_template, rental)
        pdf_bytes = render_html_to_pdf(html)
        return BytesIO(pdf_bytes)

    raise ValueError("Для ПДФ требуется разметка веб-шаблона или загруженный файл ПДФ.")


def _fill_pdf_form(template_path: str, rental: Rental) -> BytesIO:
    """
    Fill a PDF with form fields that match placeholder names (underscored),
    например customer_full_name или rental_contract_number.
    """
    field_values = {key.replace(".", "_"): value for key, value in build_placeholder_values(rental).items()}
    reader = PdfReader(template_path)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)

    if reader.get_fields():
        for page in writer.pages:
            writer.update_page_form_field_values(page, field_values)

    try:
        acroform = writer._root_object.get("/AcroForm")  # type: ignore[attr-defined]
        if acroform is not None:
            acroform.update({NameObject("/NeedAppearances"): BooleanObject(True)})
    except Exception:
        logger.debug("Не удалось пометить поток отображения ПДФ; продолжаем с исходным выводом.")

    output = BytesIO()
    writer.write(output)
    output.seek(0)
    return output
