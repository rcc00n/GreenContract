from io import BytesIO
from datetime import date
from decimal import Decimal
from unittest.mock import patch

from docx import Document
from docx.shared import Pt
from django.test import SimpleTestCase

from rentals.models import Car, ContractTemplate, Customer, Rental
from rentals.services.contract_renderer import build_placeholder_values, render_docx, render_html_template


class ContractRendererTests(SimpleTestCase):
    def _make_customer(
        self,
        full_name: str,
        driving_since: date,
        *,
        year_only: bool,
        email: str | None = None,
    ) -> Customer:
        return Customer(
            full_name=full_name,
            email=email,
            phone="+79990000000",
            license_number="12 34 567890",
            driving_since=driving_since,
            driving_since_year_only=year_only,
        )

    def _make_car(self, **overrides) -> Car:
        values = {
            "plate_number": "A001AA82",
            "make": "Hyundai",
            "model": "Solaris",
            "year": 2022,
            "daily_rate": Decimal("3000.00"),
            "color": "Белый",
            "fuel_tank_volume_liters": 50,
            "fuel_tank_cost_rub": Decimal("4200.00"),
            "security_deposit": Decimal("15000.00"),
            "rate_1_4_high": Decimal("4400.00"),
            "rate_5_14_high": Decimal("4100.00"),
            "rate_15_plus_high": Decimal("3900.00"),
            "loss_child_seat_fee": Decimal("1200.00"),
            "loss_reflective_vest_fee": Decimal("300.00"),
            "loss_registration_certificate_fee": Decimal("1500.00"),
            "loss_alloy_wheel_fee": Decimal("7000.00"),
            "loss_steel_wheel_fee": Decimal("5500.00"),
            "loss_warning_triangle_fee": Decimal("400.00"),
            "loss_radio_panel_fee": Decimal("900.00"),
            "loss_ski_mount_fee": Decimal("800.00"),
            "loss_car_keys_fee": Decimal("6500.00"),
            "loss_license_plate_fee": Decimal("2000.00"),
            "loss_hubcaps_fee": Decimal("600.00"),
            "loss_external_antenna_fee": Decimal("700.00"),
            "loss_tire_fee": Decimal("4500.00"),
            "loss_first_aid_kit_fee": Decimal("350.00"),
            "loss_jack_fee": Decimal("1000.00"),
            "loss_fire_extinguisher_fee": Decimal("500.00"),
        }
        values.update(overrides)
        return Car(**values)

    def _make_rental(
        self,
        customer: Customer,
        *,
        second_driver: Customer | None = None,
        car: Car | None = None,
    ) -> Rental:
        return Rental(
            car=car or self._make_car(),
            customer=customer,
            second_driver=second_driver,
            start_date=date(2026, 4, 2),
            end_date=date(2026, 4, 5),
            daily_rate=Decimal("3000.00"),
            total_price=Decimal("9000.00"),
        )

    def _render_docx_document(self, template_document: Document, rental: Rental) -> Document:
        template_buffer = BytesIO()
        template_document.save(template_buffer)
        template = ContractTemplate(name="test", format="docx")
        template.file = "dummy.docx"

        with patch(
            "rentals.services.contract_renderer._read_contract_template_bytes",
            return_value=template_buffer.getvalue(),
        ):
            rendered = render_docx(template, rental)

        return Document(rendered)

    def test_render_html_template_uses_year_only_for_customer_driving_since(self):
        customer = self._make_customer("Иванов Иван", date(2018, 1, 1), year_only=True)
        rental = self._make_rental(customer)
        template = ContractTemplate(
            name="test",
            format="html",
            body_html="Стаж: {{ customer.driving_since }}",
        )

        rendered = render_html_template(template, rental)

        self.assertIn("Стаж: 2018", rendered)

    def test_render_html_template_uses_year_only_for_second_driver_driving_since(self):
        customer = self._make_customer("Иванов Иван", date(2018, 1, 1), year_only=True)
        second_driver = self._make_customer("Петров Петр", date(2014, 1, 1), year_only=True)
        rental = self._make_rental(customer, second_driver=second_driver)
        template = ContractTemplate(
            name="test",
            format="html",
            body_html="Основной: {{ customer.driving_since }}; второй: {{ second_driver.driving_since }}",
        )

        rendered = render_html_template(template, rental)

        self.assertIn("Основной: 2018; второй: 2014", rendered)

    def test_render_html_template_renders_blank_for_missing_customer_email(self):
        customer = self._make_customer("Иванов Иван", date(2018, 1, 1), year_only=True, email=None)
        rental = self._make_rental(customer)
        template = ContractTemplate(
            name="test",
            format="html",
            body_html="Почта: [{{ customer.email }}]",
        )

        rendered = render_html_template(template, rental)

        self.assertIn("Почта: []", rendered)

    def test_build_placeholder_values_exposes_extended_car_fields(self):
        customer = self._make_customer("Иванов Иван", date(2018, 1, 1), year_only=True)
        rental = self._make_rental(customer)

        values = build_placeholder_values(rental)

        expected = {
            "car.rate_1_4_high": "4400",
            "car.rate_5_14_high": "4100",
            "car.rate_15_plus_high": "3900",
            "car.color": "Белый",
            "car.loss_child_seat_fee": "1200",
            "car.loss_reflective_vest_fee": "300",
            "car.loss_registration_certificate_fee": "1500",
            "car.loss_alloy_wheel_fee": "7000",
            "car.loss_steel_wheel_fee": "5500",
            "car.loss_warning_triangle_fee": "400",
            "car.loss_radio_panel_fee": "900",
            "car.loss_ski_mount_fee": "800",
            "car.loss_car_keys_fee": "6500",
            "car.loss_license_plate_fee": "2000",
            "car.loss_hubcaps_fee": "600",
            "car.fuel_tank_volume_liters": "50",
            "car.fuel_tank_cost_rub": "4200",
            "car.loss_external_antenna_fee": "700",
            "car.loss_tire_fee": "4500",
            "car.loss_first_aid_kit_fee": "350",
            "car.loss_jack_fee": "1000",
            "car.loss_fire_extinguisher_fee": "500",
            "car.security_deposit": "15000",
        }

        for key, value in expected.items():
            self.assertEqual(values[key], value, key)

    def test_build_placeholder_values_includes_all_scalar_car_fields(self):
        customer = self._make_customer("Иванов Иван", date(2018, 1, 1), year_only=True)
        rental = self._make_rental(customer)

        values = build_placeholder_values(rental)

        for field in Car._meta.concrete_fields:
            if field.is_relation or field.name == "id":
                continue
            self.assertIn(f"car.{field.name}", values, field.name)

    def test_render_docx_replaces_car_placeholders_with_flexible_whitespace(self):
        customer = self._make_customer("Иванов Иван", date(2018, 1, 1), year_only=True)
        rental = self._make_rental(customer)
        template_document = Document()
        template_document.add_paragraph(
            " | ".join(
                [
                    "{{car.rate_1_4_high}}",
                    "{{ car.rate_5_14_high }}",
                    "{{ car.rate_15_plus_high}}",
                    "{{ car.color }}",
                    "{{ car.loss_child_seat_fee }}",
                    "{{ car.loss_reflective_vest_fee }}",
                    "{{ car.loss_registration_certificate_fee }}",
                    "{{ car.loss_alloy_wheel_fee }}",
                    "{{ car.loss_steel_wheel_fee}}",
                    "{{ car.loss_warning_triangle_fee }}",
                    "{{ car.loss_radio_panel_fee }}",
                    "{{ car.loss_ski_mount_fee }}",
                    "{{ car.loss_car_keys_fee }}",
                    "{{ car.loss_license_plate_fee }}",
                    "{{ car.loss_hubcaps_fee }}",
                    "{{ car.fuel_tank_volume_liters }}",
                    "{{ car.fuel_tank_cost_rub }}",
                    "{{ car.loss_external_antenna_fee }}",
                    "{{ car.loss_tire_fee }}",
                    "{{ car.loss_first_aid_kit_fee }}",
                    "{{ car.loss_jack_fee }}",
                    "{{car.loss_fire_extinguisher_fee }}",
                    "{{ car.security_deposit  }}",
                ]
            )
        )

        rendered = self._render_docx_document(template_document, rental)
        rendered_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)

        for value in [
            "4400",
            "4100",
            "3900",
            "Белый",
            "1200",
            "300",
            "1500",
            "7000",
            "5500",
            "400",
            "900",
            "800",
            "6500",
            "2000",
            "600",
            "50",
            "4200",
            "700",
            "4500",
            "350",
            "1000",
            "500",
            "15000",
        ]:
            self.assertIn(value, rendered_text)
        self.assertNotIn("{{", rendered_text)

    def test_render_docx_preserves_run_font_sizes_while_replacing_placeholders(self):
        customer = self._make_customer("Иванов Иван", date(2018, 1, 1), year_only=True)
        rental = self._make_rental(customer)
        template_document = Document()

        paragraph = template_document.add_paragraph()
        run = paragraph.add_run("Арендатор ")
        run.font.size = Pt(7)
        run = paragraph.add_run("{{ car.color }}")
        run.font.size = Pt(7)
        run = paragraph.add_run(" ФИО")
        run.font.size = Pt(7)

        paragraph = template_document.add_paragraph()
        run = paragraph.add_run("Тариф: ")
        run.font.size = Pt(7)
        run = paragraph.add_run("{{ car.rate_1_4_high }}")
        run.font.size = Pt(9)
        run = paragraph.add_run(" руб.")
        run.font.size = Pt(7)

        paragraph = template_document.add_paragraph()
        run = paragraph.add_run("{{ car.security")
        run.font.size = Pt(7)
        run = paragraph.add_run("_deposit }}")
        run.font.size = Pt(7)

        rendered = self._render_docx_document(template_document, rental)

        self.assertEqual(rendered.paragraphs[0].text, "Арендатор Белый ФИО")
        self.assertEqual(rendered.paragraphs[1].text, "Тариф: 4400 руб.")
        self.assertEqual(rendered.paragraphs[2].text, "15000")

        paragraph0_sizes = [run.font.size.pt for run in rendered.paragraphs[0].runs if run.text]
        paragraph1_sizes = [run.font.size.pt for run in rendered.paragraphs[1].runs if run.text]
        paragraph2_sizes = [run.font.size.pt for run in rendered.paragraphs[2].runs if run.text]

        self.assertEqual(paragraph0_sizes, [7.0, 7.0, 7.0])
        self.assertEqual(paragraph1_sizes, [7.0, 9.0, 7.0])
        self.assertEqual(paragraph2_sizes, [7.0])
